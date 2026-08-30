#!/usr/bin/env python3
"""
Converter MCP Server - FastMCP Implementation
Provides tools for converting document formats using FastMCP (Streamable HTTP primary, SSE legacy).

Port allocation from ports.json only — no hardcoded ports.
FEF V3 integration preserved from original implementation.
"""
import sys
import os
import tempfile
import logging
import time
from pathlib import Path
from typing import Any
from urllib.parse import urlparse, parse_qs
import re
from contextlib import asynccontextmanager

from tools.shared.utils import is_internal_url

# ============================================================================
# Port Configuration (from ports.json only)
# ============================================================================

TOOL_NAME = "convertermcp"

try:
    from launcher.launcher_config import load_ports_config
    ports_config = load_ports_config()
    MCP_PORT = int(os.environ.get(
        "MCP_PORT",
        ports_config["assignments"]["mcp"][TOOL_NAME]
    ))
    MGMT_PORT = int(os.environ.get(
        "MCP_MGMT_PORT",
        ports_config["assignments"]["mgmt"][TOOL_NAME]
    ))
except Exception as e:
    print(f"ERROR: Failed to load ports.json: {e}", file=sys.stderr)
    sys.exit(1)

# ============================================================================
# Logging
# ============================================================================

logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(TOOL_NAME)

# ============================================================================
# Configuration
# ============================================================================

ALLOWED_ROOTS = [Path("/workspaces")]
MAX_DOCX_SIZE_MB = 20

# ============================================================================
# FEF V3 Integration
# ============================================================================

try:
    from tools.fef_integration import (
        ToolExtensionManager,
        register_common_extensions,
        setup_tool_extensions
    )
    from launcher.tool_extensions import Extension, ExtensionType
    FEF_V3_AVAILABLE = True
    logger.info("FEF V3 modules loaded successfully")
except ImportError as e:
    FEF_V3_AVAILABLE = False
    logger.warning(f"FEF V3 not available: {e}")

# ============================================================================
# Metrics
# ============================================================================

metrics = {
    "total_conversions": 0,
    "conversion_errors": 0,
    "total_conversion_time_ms": 0.0,
    "bytes_processed": 0,
}

format_usage = {}

# ============================================================================
# Helper Functions
# ============================================================================

def is_under_allowed_roots(p: Path, allowed_roots: list[Path]) -> bool:
    """Check if path p is under one of the allowed roots."""
    try:
        rp = p.resolve()
    except Exception:
        return False
    for root in allowed_roots:
        try:
            rr = root.resolve()
        except Exception:
            continue
        if rp.is_relative_to(rr):
            return True
    return False


async def download_docx_to_temp(url: str, max_size_mb: int = MAX_DOCX_SIZE_MB, headers: dict[str, str] | None = None) -> Path:
    """Download DOCX from URL to a temporary file with size checks and optional auth headers."""
    import httpx

    parsed = urlparse(url)
    if parsed.scheme not in ("http", "https"):
        raise ValueError("Only http/https URLs are allowed for DOCX download")

    if is_internal_url(url):
        raise ValueError("Cannot download DOCX from internal/private URLs")

    default_headers = {
        "User-Agent": "Mozilla/5.0 (X11; Linux x86_64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120 Safari/537.36",
        "Accept": "*/*",
    }
    merged_headers = {**default_headers, **(headers or {})}

    async with httpx.AsyncClient(timeout=60.0, follow_redirects=True, headers=merged_headers) as client:
        response = await client.get(url)
        if is_internal_url(str(response.url)):
            raise ValueError("DOCX download redirected to internal URL")
        response.raise_for_status()

        def write_temp(content_bytes: bytes) -> Path:
            size_mb = len(content_bytes) / (1024 * 1024)
            if size_mb > max_size_mb:
                raise ValueError(f"DOCX size {size_mb:.2f} MB exceeds limit of {max_size_mb} MB")
            tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
            try:
                tmp.write(content_bytes)
                tmp.flush()
            finally:
                tmp.close()
            return Path(tmp.name)

        content_type = (response.headers.get("Content-Type") or "").lower()
        if "application/vnd.openxmlformats-officedocument.wordprocessingml.document" in content_type:
            return write_temp(response.content)

        qs = parse_qs(parsed.query)
        sourcedoc_vals = qs.get("sourcedoc") or qs.get("SourceDoc") or []
        guid_raw = sourcedoc_vals[0] if sourcedoc_vals else ""
        guid_clean = guid_raw.strip().strip("{}").strip()

        m = re.search(r"(/sites/[^/]+)/", parsed.path)
        base_path = m.group(1) if m else ""
        base = f"{parsed.scheme}://{parsed.netloc}{base_path}"

        if guid_clean:
            api_url = f"{base}/_api/web/GetFileById(guid'{guid_clean}')/$value"
            api_headers = {**merged_headers, "Accept": "*/*"}
            api_resp = await client.get(api_url, headers=api_headers)
            if api_resp.status_code == 200 and (api_resp.headers.get("Content-Type", "").lower().startswith("application/") or api_resp.content):
                return write_temp(api_resp.content)

        raise ValueError(f"Expected DOCX but received Content-Type: '{content_type}'. For SharePoint Doc.aspx, use the underlying download endpoint or ensure sourcedoc GUID is accessible via _api.")


def extract_docx_text(docx_path: Path) -> str:
    """Extract plain text from DOCX preserving paragraphs and table content."""
    from docx import Document
    doc = Document(str(docx_path))
    lines: list[str] = []
    for para in doc.paragraphs:
        lines.append(para.text)
    for table in doc.tables:
        lines.append("")
        for row in table.rows:
            row_text = "\t".join(cell.text.strip() for cell in row.cells)
            if row_text.strip():
                lines.append(row_text)
    return "\n".join(lines)


# ============================================================================
# FastMCP Instance (via shared factory — DualHeaderVerifier auth)
# ============================================================================

from tools.shared.server_factory import create_fastmcp_server, DEFAULT_HOST

mcp = create_fastmcp_server(TOOL_NAME)


@mcp.tool()
async def convert_docx_to_text(source: str, output_path: str | None = None, headers: dict[str, str] | None = None) -> str:
    """Converts a DOCX file (local path or URL) to plain text; writes to output_path if provided, otherwise returns text.

    Args:
        source: Local path or URL to a DOCX file
        output_path: Optional local path to write the extracted text
        headers: Optional dict of HTTP headers for URL downloads (e.g., auth)
    """
    temp_path: Path | None = None
    start_time = time.perf_counter()
    success = False

    try:
        if source.startswith("http://") or source.startswith("https://"):
            logger.info(f"Downloading DOCX from URL: {source}")
            temp_path = await download_docx_to_temp(source, headers=headers)
            docx_path = temp_path
        else:
            docx_path = Path(source)
            if not docx_path.exists():
                elapsed_ms = (time.perf_counter() - start_time) * 1000
                if fef_manager is not None:
                    fef_manager.metrics.record_request(
                        endpoint="tools/call", tool_name="convert_docx_to_text",
                        success=False, duration_ms=elapsed_ms
                    )
                return f"Error: File not found: {docx_path}"
            if not is_under_allowed_roots(docx_path, ALLOWED_ROOTS):
                elapsed_ms = (time.perf_counter() - start_time) * 1000
                if fef_manager is not None:
                    fef_manager.metrics.record_request(
                        endpoint="tools/call", tool_name="convert_docx_to_text",
                        success=False, duration_ms=elapsed_ms
                    )
                return f"Error: Path not allowed: {docx_path}"

        logger.info(f"Converting DOCX to text: {docx_path}")
        text = extract_docx_text(docx_path)

        if output_path:
            out_path = Path(output_path)
            if not is_under_allowed_roots(out_path, ALLOWED_ROOTS):
                elapsed_ms = (time.perf_counter() - start_time) * 1000
                metrics["conversion_errors"] += 1
                if fef_manager is not None:
                    fef_manager.metrics.record_request(
                        endpoint="tools/call", tool_name="convert_docx_to_text",
                        success=False, duration_ms=elapsed_ms
                    )
                return f"Error: Output path not allowed: {out_path}"
            out_path.parent.mkdir(parents=True, exist_ok=True)
            with Path(out_path).open("w", encoding="utf-8") as f:
                f.write(text)
            docx_size = docx_path.stat().st_size if docx_path.exists() else 0
            txt_size = out_path.stat().st_size
            elapsed_ms = (time.perf_counter() - start_time) * 1000
            metrics["total_conversions"] += 1
            metrics["total_conversion_time_ms"] += elapsed_ms
            metrics["bytes_processed"] += docx_size
            if fef_manager is not None:
                fef_manager.metrics.record_request(
                    endpoint="tools/call", tool_name="convert_docx_to_text",
                    success=True, duration_ms=elapsed_ms
                )
            msg = f"Success: {docx_size:,} bytes DOCX converted to {txt_size:,} bytes TXT at {out_path}"
            logger.info(msg)
            success = True
            return msg

        elapsed_ms = (time.perf_counter() - start_time) * 1000
        metrics["total_conversions"] += 1
        metrics["total_conversion_time_ms"] += elapsed_ms
        metrics["bytes_processed"] += docx_path.stat().st_size if docx_path.exists() else 0
        success = True
        return text

    except ImportError as e:
        logger.error(f"Dependency error: {e}")
        metrics["conversion_errors"] += 1
        elapsed_ms = (time.perf_counter() - start_time) * 1000
        if fef_manager is not None:
            fef_manager.metrics.record_request(
                endpoint="tools/call", tool_name="convert_docx_to_text",
                success=False, duration_ms=elapsed_ms
            )
        return f"Error: {str(e)}"
    except Exception as e:
        logger.error(f"Error converting DOCX: {e}")
        metrics["conversion_errors"] += 1
        elapsed_ms = (time.perf_counter() - start_time) * 1000
        if fef_manager is not None:
            fef_manager.metrics.record_request(
                endpoint="tools/call", tool_name="convert_docx_to_text",
                success=False, duration_ms=elapsed_ms
            )
        return f"Error: {str(e)}"
    finally:
        if temp_path:
            try:
                temp_path.unlink(missing_ok=True)
            except Exception:
                pass


# ============================================================================
# FEF V3 Data Sources
# ============================================================================

def get_conversion_stats(params: dict[str, Any]) -> dict[str, Any]:
    avg_time = (
        metrics["total_conversion_time_ms"] / metrics["total_conversions"]
        if metrics["total_conversions"] > 0 else 0.0
    )
    return {
        "total_conversions": metrics["total_conversions"],
        "conversion_errors": metrics["conversion_errors"],
        "avg_conversion_time_ms": round(avg_time, 2),
        "bytes_processed": metrics["bytes_processed"]
    }


def get_format_usage(params: dict[str, Any]) -> dict[str, Any]:
    return {
        "formats": format_usage,
        "total": sum(format_usage.values())
    }


def get_conversion_queue(params: dict[str, Any]) -> dict[str, Any]:
    return {
        "pending_conversions": metrics.get("pending", 0),
        "active_conversions": metrics.get("active", 0),
        "completed_today": metrics["total_conversions"]
    }


def get_storage_usage(params: dict[str, Any]) -> dict[str, Any]:
    return {
        "total_bytes_processed": metrics["bytes_processed"],
        "estimated_disk_usage_mb": round(metrics["bytes_processed"] / (1024 * 1024), 2)
    }


# ============================================================================
# FEF V3 Extensions Setup
# ============================================================================

fef_manager = None
fef_registry = None
fef_http_server = None
fef_setup_done = False


def setup_extensions(registry=None) -> None:
    global fef_manager, fef_registry, fef_http_server, fef_setup_done

    if fef_setup_done:
        return

    if not FEF_V3_AVAILABLE:
        fef_setup_done = True
        return

    mgmt_port = int(os.environ.get("MCP_MGMT_PORT", MGMT_PORT))

    custom_extensions = [
        Extension(name="conversion_stats", ext_type=ExtensionType.DATA_SOURCE,
            schema={"input": {"type": "object", "properties": {}}, "output": {"type": "object", "properties": {"total_conversions": {"type": "integer"}, "conversion_errors": {"type": "integer"}, "avg_conversion_time_ms": {"type": "number"}}}},
            handler=get_conversion_stats, metadata={"description": "Document conversion statistics", "category": "metrics"}),
        Extension(name="format_usage", ext_type=ExtensionType.DATA_SOURCE,
            schema={"input": {"type": "object", "properties": {}}, "output": {"type": "object", "properties": {"formats": {"type": "object"}, "total": {"type": "integer"}}}},
            handler=get_format_usage, metadata={"description": "Format usage statistics", "category": "metrics"}),
        Extension(name="conversion_queue", ext_type=ExtensionType.DATA_SOURCE,
            schema={"input": {"type": "object", "properties": {}}, "output": {"type": "object", "properties": {"pending_conversions": {"type": "integer"}, "active_conversions": {"type": "integer"}, "completed_today": {"type": "integer"}}}},
            handler=get_conversion_queue, metadata={"description": "Conversion queue status", "category": "metrics"}),
        Extension(name="storage_usage", ext_type=ExtensionType.DATA_SOURCE,
            schema={"input": {"type": "object", "properties": {}}, "output": {"type": "object", "properties": {"total_bytes_processed": {"type": "integer"}, "estimated_disk_usage_mb": {"type": "number"}}}},
            handler=get_storage_usage, metadata={"description": "Storage usage by output", "category": "metrics"}),
    ]

    if registry is not None:
        fef_registry = registry
        fef_manager = ToolExtensionManager(TOOL_NAME)
        register_common_extensions(TOOL_NAME, fef_registry, fef_manager)
        for ext in custom_extensions:
            fef_registry.register(TOOL_NAME, ext)
        fef_http_server = None
        logger.info(f"[{TOOL_NAME}] FEF V3 registered with launcher's registry")
    else:
        fef_manager, fef_registry, fef_http_server = setup_tool_extensions(
            tool_name=TOOL_NAME, mgmt_port=mgmt_port, custom_extensions=custom_extensions)
        logger.info(f"[{TOOL_NAME}] FEF V3 standalone mode on port {mgmt_port}")

    fef_setup_done = True


# ============================================================================
# Lifespan
# ============================================================================

@asynccontextmanager
async def lifespan(app):
    global fef_http_server

    logger.info(f"{TOOL_NAME} FastMCP server starting on port {MCP_PORT}...")

    if not fef_setup_done:
        setup_extensions(registry=None)

    if FEF_V3_AVAILABLE and fef_http_server:
        try:
            await fef_http_server.start()
        except Exception as e:
            logger.warning(f"Failed to start FEF V3 management server: {e}")

    yield

    logger.info(f"{TOOL_NAME} FastMCP server shutting down...")
    if fef_http_server:
        try:
            await fef_http_server.stop()
        except Exception:
            pass


# ============================================================================
# App Export
# Transports: streamable-http (default, /mcp) or sse (legacy compat, /sse + /messages)
# ============================================================================

from tools.shared.server_factory import get_transport_app

app = get_transport_app(mcp)

__all__ = ["app", "setup_extensions", "mcp"]


# ============================================================================
# Main Entry Point
# ============================================================================

if __name__ == "__main__":
    import uvicorn

    transport = os.environ.get("MCP_TRANSPORT", "streamable-http").lower()
    logger.info(f"Starting {TOOL_NAME} FastMCP server (transport: {transport})")
    logger.info(f"  MCP port: {MCP_PORT}")
    if transport == "sse":
        logger.info(f"  SSE endpoint: http://localhost:{MCP_PORT}/sse")
    else:
        logger.info(f"  Streamable HTTP: http://localhost:{MCP_PORT}/mcp")
    if FEF_V3_AVAILABLE:
        logger.info(f"  FEF V3 mgmt: http://localhost:{MGMT_PORT}")

    uvicorn.run(app, host=DEFAULT_HOST, port=MCP_PORT, log_level="info", lifespan="on")
