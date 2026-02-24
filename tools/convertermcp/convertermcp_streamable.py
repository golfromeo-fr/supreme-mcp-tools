#!/usr/bin/env python3
"""
Converter MCP Server - Streamable HTTP Transport
Provides tools for converting document formats (DOCX to text, etc.)
"""
import sys
import os
import tempfile
import logging
from pathlib import Path
from typing import Any, AsyncGenerator, Dict, Optional
from contextlib import asynccontextmanager

# Check for required dependencies before importing
try:
    from fastapi import FastAPI, Request, Response
    from fastapi.responses import StreamingResponse
    import uvicorn
    import httpx
    from urllib.parse import urlparse, parse_qs
    import re
except ImportError as e:
    print(f"ERROR: Missing required dependency: {e}", file=sys.stderr)
    print("Please make sure the virtual environment is activated and all dependencies are installed.", file=sys.stderr)
    print("Run: pip install -r requirements.txt", file=sys.stderr)
    sys.exit(1)

# Add parent directories to path for importing StreamableHttpTransportBase
# The supreme-mcp-tools directory (parent of tools and launcher) needs to be in the path
# Script is at: tools/convertermcp/convertermcp_streamable.py
# supreme-mcp-tools is at: . (relative path)
supreme_mcp_tools_dir = os.path.abspath(os.path.join(os.path.dirname(__file__), "../.."))
if supreme_mcp_tools_dir not in sys.path:
    sys.path.insert(0, supreme_mcp_tools_dir)

try:
    from launcher.streamable_http.streamable_http_base import (
        StreamableHttpTransportBase,
        StreamableHttpConfig,
    )
except ImportError as e:
    print(f"ERROR: Cannot import StreamableHttpTransportBase: {e}", file=sys.stderr)
    print(f"Script location: {__file__}", file=sys.stderr)
    print(f"supreme_mcp_tools_dir: {supreme_mcp_tools_dir}", file=sys.stderr)
    print(f"Python path: {sys.path}", file=sys.stderr)
    print("Please ensure the launcher/streamable_http module is available.", file=sys.stderr)
    print("Try running from the supreme-mcp-tools directory: python tools/convertermcp/convertermcp_streamable.py", file=sys.stderr)
    sys.exit(1)

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s'
)

# Add protocol version compatibility
SUPPORTED_PROTOCOL_VERSIONS = ["2024-11-05", "2025-11-25"]
logger = logging.getLogger("convertermcp_streamable")

# ============================================================================
# Configuration
# ============================================================================

# Allowed roots for local file read/write operations for DOCX conversion
ALLOWED_ROOTS = [
    Path("/workspaces")
]

# Maximum allowed DOCX size when downloading from URL (in MB)
MAX_DOCX_SIZE_MB = 20

# ============================================================================
# Helper Functions
# ============================================================================

def is_under_allowed_roots(p: Path, allowed_roots: list[Path]) -> bool:
    """Check if path p is under one of the allowed roots."""
    try:
        rp = p.resolve()
    except Exception:
        return False
    for root in allowed_roots:
        try:
            rr = root.resolve()
        except Exception:
            continue
        if str(rp).startswith(str(rr)):
            return True
    return False


async def download_docx_to_temp(url: str, max_size_mb: int = MAX_DOCX_SIZE_MB, headers: dict | None = None) -> Path:
    """Download DOCX from URL to a temporary file with size checks and optional auth headers.

    SharePoint Doc.aspx fallback:
    - If the initial request returns HTML instead of a DOCX, try the SharePoint REST API using the sourcedoc GUID:
      GET {base}/_api/web/GetFileById(guid'{GUID}')/$value
    - base is inferred from '/sites/{site}/' in the URL path when present, else scheme://host
    """
    parsed = urlparse(url)
    if parsed.scheme not in ("http", "https"):
        raise ValueError("Only http/https URLs are allowed for DOCX download")

    default_headers = {
        "User-Agent": "Mozilla/5.0 (X11; Linux x86_64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120 Safari/537.36",
        "Accept": "*/*",
    }
    merged_headers = {**default_headers, **(headers or {})}

    async with httpx.AsyncClient(timeout=60.0, follow_redirects=True, headers=merged_headers) as client:
        # 1) First attempt: GET the provided URL
        response = await client.get(url)
        response.raise_for_status()

        def write_temp(content_bytes: bytes) -> Path:
            size_mb = len(content_bytes) / (1024 * 1024)
            if size_mb > max_size_mb:
                raise ValueError(f"DOCX size {size_mb:.2f} MB exceeds limit of {max_size_mb} MB")
            tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
            try:
                tmp.write(content_bytes)
                tmp.flush()
            finally:
                tmp.close()
            return Path(tmp.name)

        content_type = (response.headers.get("Content-Type") or "").lower()
        if "application/vnd.openxmlformats-officedocument.wordprocessingml.document" in content_type:
            # Looks like a DOCX response
            return write_temp(response.content)

        # 2) Fallback for SharePoint Doc.aspx: use _api with sourcedoc GUID if present
        # Extract sourcedoc GUID from query
        # sourcedoc={%7BGUID%7D} or {GUID}
        qs = parse_qs(parsed.query)
        sourcedoc_vals = qs.get("sourcedoc") or qs.get("SourceDoc") or []
        guid_raw = sourcedoc_vals[0] if sourcedoc_vals else ""
        guid_clean = guid_raw.strip().strip("{}").strip()

        # Build base = scheme://host[/sites/{site}]
        # Try to capture '/sites/XYZ/' from the path
        m = re.search(r"(/sites/[^/]+)/", parsed.path)
        base_path = m.group(1) if m else ""
        base = f"{parsed.scheme}://{parsed.netloc}{base_path}"

        if guid_clean:
            api_url = f"{base}/_api/web/GetFileById(guid'{guid_clean}')/$value"
            # Accept a binary stream
            api_headers = {**merged_headers, "Accept": "*/*"}
            api_resp = await client.get(api_url, headers=api_headers)
            if api_resp.status_code == 200 and (api_resp.headers.get("Content-Type", "").lower().startswith("application/") or api_resp.content):
                return write_temp(api_resp.content)

        # 3) If still not a DOCX, raise a diagnostic error with content-type
        raise ValueError(f"Expected DOCX but received Content-Type: '{content_type}'. For SharePoint Doc.aspx, use the underlying download endpoint or ensure sourcedoc GUID is accessible via _api.")


def extract_docx_text(docx_path: Path) -> str:
    """Extract plain text from DOCX preserving paragraphs and table content."""
    try:
        from docx import Document
    except ImportError:
        raise ImportError("python-docx not installed. Install with: pip install python-docx")
    doc = Document(str(docx_path))
    lines: list[str] = []
    # Paragraphs
    for para in doc.paragraphs:
        lines.append(para.text)
    # Tables (tab-separated rows with separators)
    for table in doc.tables:
        lines.append("")  # separator before table
        for row in table.rows:
            row_text = "\t".join(cell.text.strip() for cell in row.cells)
            if row_text.strip():
                lines.append(row_text)
        lines.append("")  # separator after table
    return "\n".join(lines)


# ============================================================================
# Converter MCP Streamable HTTP Transport Implementation
# ============================================================================

class ConverterMCPStreamableHttp(StreamableHttpTransportBase):
    """
    Converter MCP server implementation using Streamable HTTP transport.
    
    This class provides document conversion tools (DOCX to text) using
    the Streamable HTTP transport with JSON-RPC framing.
    """
    
    def __init__(self):
        """Initialize the Converter MCP Streamable HTTP server."""
        config = StreamableHttpConfig(
            endpoint="/mcp",
            framing_format="newline-delimited",
            request_timeout=60.0,  # Longer timeout for file downloads
        )
        super().__init__("convertermcp", config)
        logger.info("Converter MCP Streamable HTTP transport initialized")
    
    async def _handle_initialize(self, params, session):
        """Handle initialize request - only tools are supported."""
        protocol_version = params.get("protocolVersion", "2024-11-05")
        # Support both old and new protocol versions
        if protocol_version not in SUPPORTED_PROTOCOL_VERSIONS:
            logger.warning(f"Client sent unsupported protocol version: {protocol_version}, using 2024-11-05")
            protocol_version = "2024-11-05"
        
        # Return server capabilities - only tools are supported (matching original convertermcp)
        return {
            "jsonrpc": "2.0",
            "result": {
                "protocolVersion": protocol_version,
                "capabilities": {
                    "tools": {},  # Tools are supported
                    # resources and prompts are not included, indicating they're not supported
                },
                "serverInfo": {
                    "name": self.server_name,
                    "version": "1.0.0",
                },
            },
        }
    
    async def _handle_tools_list(
        self,
        params: Dict[str, Any],
        session: Dict[str, Any]
    ) -> Dict[str, Any]:
        """Handle tools/list request."""
        tools = [
            {
                "name": "convert_docx_to_text",
                "description": "Converts a DOCX file (local path or URL) to plain text; writes to output_path if provided, otherwise returns text.",
                "inputSchema": {
                    "type": "object",
                    "required": ["source"],
                    "properties": {
                        "source": {
                            "type": "string",
                            "description": "Local .docx path or http/https URL"
                        },
                        "output_path": {
                            "type": "string",
                            "description": "Optional path to write the output text file"
                        },
                        "headers": {
                            "type": "object",
                            "description": "Optional HTTP headers for authenticated URL download (e.g., Authorization, Cookie)"
                        }
                    }
                }
            }
        ]
        
        return {
            "jsonrpc": "2.0",
            "result": {
                "tools": tools,
            },
        }
    
    async def _handle_tool_call(
        self,
        params: Dict[str, Any],
        session: Dict[str, Any],
        request_id: Any
    ) -> AsyncGenerator[Dict[str, Any], None]:
        """Handle tools/call request."""
        tool_name = params.get("name")
        arguments = params.get("arguments", {})
        
        logger.info(f"Tool called: {tool_name} with arguments: {arguments}")
        
        try:
            if tool_name == "convert_docx_to_text":
                async for response in self._handle_convert_docx_to_text(arguments, request_id):
                    yield response
            else:
                yield {
                    "jsonrpc": "2.0",
                    "id": request_id,
                    "error": {
                        "code": -32601,
                        "message": f"Unknown tool: {tool_name}"
                    }
                }
        
        except ValueError as e:
            logger.error(f"Value error in tool call '{tool_name}': {e}")
            yield {
                "jsonrpc": "2.0",
                "id": request_id,
                "error": {
                    "code": -32602,
                    "message": "Invalid params",
                    "data": str(e)
                }
            }
        
        except httpx.HTTPError as e:
            logger.error(f"HTTP error in tool call '{tool_name}': {e}")
            yield {
                "jsonrpc": "2.0",
                "id": request_id,
                "error": {
                    "code": -32603,
                    "message": "HTTP error",
                    "data": str(e)
                }
            }
        
        except Exception as e:
            logger.error(f"Error in tool call '{tool_name}': {e}")
            yield {
                "jsonrpc": "2.0",
                "id": request_id,
                "error": {
                    "code": -32603,
                    "message": "Internal error",
                    "data": str(e)
                }
            }
    
    async def _handle_convert_docx_to_text(
        self,
        arguments: Dict[str, Any],
        request_id: Any
    ) -> AsyncGenerator[Dict[str, Any], None]:
        """Handler for converting DOCX to text from local path or URL; optionally writes to output_path."""
        source = arguments.get("source")
        output_path_arg = arguments.get("output_path")
        headers_arg = arguments.get("headers")
        
        # Validate headers
        if headers_arg is not None and not isinstance(headers_arg, dict):
            logger.error("Invalid headers: must be an object of string key/value pairs")
            yield {
                "jsonrpc": "2.0",
                "id": request_id,
                "error": {
                    "code": -32602,
                    "message": "Invalid params",
                    "data": "headers must be an object of string key/value pairs"
                }
            }
            return
        
        if not source or not isinstance(source, str):
            logger.error("Missing required parameter: source")
            yield {
                "jsonrpc": "2.0",
                "id": request_id,
                "error": {
                    "code": -32602,
                    "message": "Invalid params",
                    "data": "Missing required parameter: source"
                }
            }
            return
        
        temp_path: Path | None = None
        try:
            # Determine source type
            if source.startswith("http://") or source.startswith("https://"):
                logger.info(f"Downloading DOCX from URL: {source}")
                temp_path = await download_docx_to_temp(source, headers=headers_arg)
                docx_path = temp_path
            else:
                docx_path = Path(source)
                if not docx_path.exists():
                    logger.error(f"File not found: {docx_path}")
                    yield {
                        "jsonrpc": "2.0",
                        "id": request_id,
                        "error": {
                            "code": -32602,
                            "message": "Invalid params",
                            "data": f"File not found: {docx_path}"
                        }
                    }
                    return
                if not is_under_allowed_roots(docx_path, ALLOWED_ROOTS):
                    logger.error(f"Path not allowed: {docx_path}")
                    yield {
                        "jsonrpc": "2.0",
                        "id": request_id,
                        "error": {
                            "code": -32602,
                            "message": "Invalid params",
                            "data": f"Path not allowed: {docx_path}"
                        }
                    }
                    return
            
            # Convert to text
            logger.info(f"Converting DOCX to text: {docx_path}")
            text = extract_docx_text(docx_path)
            
            # Prefer writing to file if output_path provided
            if output_path_arg:
                out_path = Path(output_path_arg)
                if not is_under_allowed_roots(out_path, ALLOWED_ROOTS):
                    logger.error(f"Output path not allowed: {out_path}")
                    yield {
                        "jsonrpc": "2.0",
                        "id": request_id,
                        "error": {
                            "code": -32602,
                            "message": "Invalid params",
                            "data": f"Output path not allowed: {out_path}"
                        }
                    }
                    return
                out_path.parent.mkdir(parents=True, exist_ok=True)
                with open(out_path, "w", encoding="utf-8") as f:
                    f.write(text)
                docx_size = docx_path.stat().st_size if docx_path.exists() else 0
                txt_size = out_path.stat().st_size
                msg = f"Success: {docx_size:,} bytes DOCX converted to {txt_size:,} bytes TXT at {out_path}"
                logger.info(msg)
                yield {
                    "jsonrpc": "2.0",
                    "id": request_id,
                    "result": {
                        "content": [
                            {
                                "type": "text",
                                "text": msg
                            }
                        ]
                    }
                }
            else:
                # Otherwise return text to LLM
                logger.info("Returning extracted text to the LLM")
                yield {
                    "jsonrpc": "2.0",
                    "id": request_id,
                    "result": {
                        "content": [
                            {
                                "type": "text",
                                "text": text
                            }
                        ]
                    }
                }
        
        except ImportError as e:
            logger.error(f"Dependency error: {e}")
            yield {
                "jsonrpc": "2.0",
                "id": request_id,
                "error": {
                    "code": -32603,
                    "message": "Internal error",
                    "data": str(e)
                }
            }
        
        except Exception as e:
            logger.error(f"Error converting DOCX: {e}")
            yield {
                "jsonrpc": "2.0",
                "id": request_id,
                "error": {
                    "code": -32603,
                    "message": "Internal error",
                    "data": str(e)
                }
            }
        
        finally:
            # Cleanup temporary file if used
            if temp_path:
                try:
                    temp_path.unlink(missing_ok=True)
                except Exception:
                    pass


# ============================================================================
# FastAPI Application
# ============================================================================

# Create the transport instance
transport = ConverterMCPStreamableHttp()

# Lifespan context manager for startup/shutdown events
@asynccontextmanager
async def lifespan(app: FastAPI):
    """Handle application lifespan events."""
    # Startup
    logger.info("Converter MCP Streamable HTTP server starting up...")
    yield
    # Shutdown
    logger.info("Converter MCP Streamable HTTP server shutting down...")
    await transport.cleanup_sessions()

# Create FastAPI application with lifespan
app = FastAPI(
    title="Converter MCP Streamable HTTP Server",
    description="Document conversion tools using Streamable HTTP transport",
    version="1.0.0",
    lifespan=lifespan
)


@app.get("/")
async def root():
    """Root endpoint with server information."""
    return {
        "name": "convertermcp",
        "version": "1.0.0",
        "transport": "streamable-http",
        "endpoint": "/mcp",
        "tools": ["convert_docx_to_text"]
    }


@app.get("/health")
async def health():
    """Health check endpoint."""
    return {
        "status": "healthy",
        "active_sessions": transport.get_session_count()
    }


@app.post("/mcp")
async def handle_mcp_request(request: Request):
    """
    Handle Streamable HTTP MCP requests.
    
    This endpoint accepts JSON-RPC requests with proper framing and returns
    responses with the same framing format.
    """
    # Read request body
    body = await request.body()
    logger.debug(f"Received request body: {body}")
    
    # Extract headers
    headers = dict(request.headers)
    session_id = headers.get("Mcp-Session-Id")
    logger.debug(f"Request headers: {headers}, session_id: {session_id}")
    
    # Parse request data (expecting newline-delimited JSON)
    try:
        import json
        request_data = json.loads(body.decode("utf-8").strip())
        logger.info(f"Processing JSON-RPC request: method={request_data.get('method')}, id={request_data.get('id')}")
    except json.JSONDecodeError as e:
        logger.error(f"Failed to parse request: {e}")
        return Response(
            content=json.dumps({
                "jsonrpc": "2.0",
                "error": {
                    "code": -32700,
                    "message": "Parse error",
                    "data": str(e)
                }
            }),
            status_code=400,
            media_type="application/json"
        )
    
    # Process the request
    async def response_generator():
        async for response in transport.handle_request(request_data, headers, session_id):
            # Format response based on framing configuration
            logger.debug(f"Generating response: {response}")
            if transport.config.framing_format == "newline-delimited":
                yield (json.dumps(response) + "\n").encode("utf-8")
            else:
                yield json.dumps(response).encode("utf-8")
    
    return StreamingResponse(
        response_generator(),
        media_type="application/json",
        headers={
            "Content-Type": "application/json",
            "X-MCP-Transport": "streamable-http",
            "X-MCP-Framing": transport.config.framing_format,
        }
    )


# ============================================================================
# Main Entry Point
# ============================================================================

if __name__ == "__main__":
    import argparse
    
    parser = argparse.ArgumentParser(description="Converter MCP Streamable HTTP Server")
    parser.add_argument(
        "--host",
        default="0.0.0.0",
        help="Host to bind to (default: 0.0.0.0)"
    )
    parser.add_argument(
        "--port",
        type=int,
        default=8003,
        help="Port to bind to (default: 8003)"
    )
    parser.add_argument(
        "--log-level",
        default="info",
        choices=["debug", "info", "warning", "error"],
        help="Log level (default: info)"
    )
    
    args = parser.parse_args()
    
    # Set log level
    log_level = getattr(logging, args.log_level.upper())
    logging.getLogger("convertermcp_streamable").setLevel(log_level)
    
    logger.info(f"Starting Converter MCP Streamable HTTP Server on http://{args.host}:{args.port}")
    
    try:
        uvicorn.run(
            app,
            host=args.host,
            port=args.port,
            log_level=args.log_level,
        )
    except KeyboardInterrupt:
        logger.info("Server shutting down gracefully...")
        sys.exit(0)
    except Exception as e:
        logger.error(f"Server error: {e}")
        sys.exit(1)


"""
VSCode Configuration Example (Streamable HTTP):

{
  "mcpServers": {
    "convertermcp": {
      "type": "streamable-http",
      "url": "http://localhost:8003/mcp",
      "headers": {
        "Content-Type": "application/json"
      },
      "framing": "newline-delimited"
    }
  }
}
"""
