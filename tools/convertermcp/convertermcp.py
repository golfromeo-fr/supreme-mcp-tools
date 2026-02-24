#!/usr/bin/env python3
"""
Converter MCP Server - Document Conversion Tools
Provides tools for converting document formats (DOCX to text, etc.)
"""
import sys
import os
import tempfile
import logging
from pathlib import Path
from typing import Optional

# Check for required dependencies before importing
try:
    import anyio
    import mcp.types as types
    from mcp.server.lowlevel import Server
    from mcp.server.sse import SseServerTransport
    from starlette.applications import Starlette
    from starlette.routing import Route, Mount
    import httpx
    from urllib.parse import urlparse, parse_qs
    import re
except ImportError as e:
    print(f"ERROR: Missing required dependency: {e}", file=sys.stderr)
    print("Please make sure the virtual environment is activated and all dependencies are installed.", file=sys.stderr)
    print("Run: pip install -r requirements.txt", file=sys.stderr)
    sys.exit(1)

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger("convertermcp")

# ============================================================================
# Configuration
# ============================================================================

# Allowed roots for local file read/write operations for DOCX conversion
ALLOWED_ROOTS = [
    Path("/workspaces")
]

# Maximum allowed DOCX size when downloading from URL (in MB)
MAX_DOCX_SIZE_MB = 20

# ============================================================================
# Server Initialization
# ============================================================================

# Verify server components
logger.info("Initializing Converter MCP Server...")

try:
    server = Server("convertermcp")
    sse_transport = SseServerTransport("/messages/")
    logger.info("Server components initialized successfully")
except Exception as e:
    logger.error(f"Failed to initialize server components: {e}")
    sys.exit(1)


# ============================================================================
# Helper Functions
# ============================================================================

def is_under_allowed_roots(p: Path, allowed_roots: list[Path]) -> bool:
    """Check if path p is under one of the allowed roots."""
    try:
        rp = p.resolve()
    except Exception:
        return False
    for root in allowed_roots:
        try:
            rr = root.resolve()
        except Exception:
            continue
        if str(rp).startswith(str(rr)):
            return True
    return False


async def download_docx_to_temp(url: str, max_size_mb: int = MAX_DOCX_SIZE_MB, headers: dict | None = None) -> Path:
    """Download DOCX from URL to a temporary file with size checks and optional auth headers.

    SharePoint Doc.aspx fallback:
    - If the initial request returns HTML instead of a DOCX, try the SharePoint REST API using the sourcedoc GUID:
      GET {base}/_api/web/GetFileById(guid'{GUID}')/$value
    - base is inferred from '/sites/{site}/' in the URL path when present, else scheme://host
    """
    parsed = urlparse(url)
    if parsed.scheme not in ("http", "https"):
        raise ValueError("Only http/https URLs are allowed for DOCX download")

    default_headers = {
        "User-Agent": "Mozilla/5.0 (X11; Linux x86_64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120 Safari/537.36",
        "Accept": "*/*",
    }
    merged_headers = {**default_headers, **(headers or {})}

    async with httpx.AsyncClient(timeout=60.0, follow_redirects=True, headers=merged_headers) as client:
        # 1) First attempt: GET the provided URL
        response = await client.get(url)
        response.raise_for_status()

        def write_temp(content_bytes: bytes) -> Path:
            size_mb = len(content_bytes) / (1024 * 1024)
            if size_mb > max_size_mb:
                raise ValueError(f"DOCX size {size_mb:.2f} MB exceeds limit of {max_size_mb} MB")
            tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
            try:
                tmp.write(content_bytes)
                tmp.flush()
            finally:
                tmp.close()
            return Path(tmp.name)

        content_type = (response.headers.get("Content-Type") or "").lower()
        if "application/vnd.openxmlformats-officedocument.wordprocessingml.document" in content_type:
            # Looks like a DOCX response
            return write_temp(response.content)

        # 2) Fallback for SharePoint Doc.aspx: use _api with sourcedoc GUID if present
        # Extract sourcedoc GUID from query
        # sourcedoc={%7BGUID%7D} or {GUID}
        qs = parse_qs(parsed.query)
        sourcedoc_vals = qs.get("sourcedoc") or qs.get("SourceDoc") or []
        guid_raw = sourcedoc_vals[0] if sourcedoc_vals else ""
        guid_clean = guid_raw.strip().strip("{}").strip()

        # Build base = scheme://host[/sites/{site}]
        # Try to capture '/sites/XYZ/' from the path
        m = re.search(r"(/sites/[^/]+)/", parsed.path)
        base_path = m.group(1) if m else ""
        base = f"{parsed.scheme}://{parsed.netloc}{base_path}"

        if guid_clean:
            api_url = f"{base}/_api/web/GetFileById(guid'{guid_clean}')/$value"
            # Accept a binary stream
            api_headers = {**merged_headers, "Accept": "*/*"}
            api_resp = await client.get(api_url, headers=api_headers)
            if api_resp.status_code == 200 and (api_resp.headers.get("Content-Type", "").lower().startswith("application/") or api_resp.content):
                return write_temp(api_resp.content)

        # 3) If still not a DOCX, raise a diagnostic error with content-type
        raise ValueError(f"Expected DOCX but received Content-Type: '{content_type}'. For SharePoint Doc.aspx, use the underlying download endpoint or ensure sourcedoc GUID is accessible via _api.")


def extract_docx_text(docx_path: Path) -> str:
    """Extract plain text from DOCX preserving paragraphs and table content."""
    try:
        from docx import Document
    except ImportError:
        raise ImportError("python-docx not installed. Install with: pip install python-docx")
    doc = Document(str(docx_path))
    lines: list[str] = []
    # Paragraphs
    for para in doc.paragraphs:
        lines.append(para.text)
    # Tables (tab-separated rows with separators)
    for table in doc.tables:
        lines.append("")  # separator before table
        for row in table.rows:
            row_text = "\t".join(cell.text.strip() for cell in row.cells)
            if row_text.strip():
                lines.append(row_text)
        lines.append("")  # separator after table
    return "\n".join(lines)


# ============================================================================
# SSE Handler
# ============================================================================

async def handle_sse(request):
    from starlette.responses import Response
    
    async with sse_transport.connect_sse(request.scope, request.receive, request._send) as streams:
        await server.run(streams[0], streams[1], server.create_initialization_options())
    return Response()


# ============================================================================
# Tool: Convert DOCX to Text
# ============================================================================

async def handle_convert_docx_to_text_tool(arguments: dict) -> list[types.TextContent]:
    """Handler for converting DOCX to text from local path or URL; optionally writes to output_path."""
    source = arguments.get("source")
    output_path_arg = arguments.get("output_path")
    headers_arg = arguments.get("headers")
    if headers_arg is not None and not isinstance(headers_arg, dict):
        logger.error("Invalid headers: must be an object of string key/value pairs")
        return [types.TextContent(type="text", text="Error: headers must be an object of string key/value pairs")]
    if not source or not isinstance(source, str):
        logger.error("Missing required parameter: source")
        return [types.TextContent(type="text", text="Missing required parameter: source")]
    temp_path: Path | None = None
    try:
        # Determine source type
        if source.startswith("http://") or source.startswith("https://"):
            logger.info(f"Downloading DOCX from URL: {source}")
            temp_path = await download_docx_to_temp(source, headers=headers_arg)
            docx_path = temp_path
        else:
            docx_path = Path(source)
            if not docx_path.exists():
                logger.error(f"File not found: {docx_path}")
                return [types.TextContent(type="text", text=f"Error: File not found: {docx_path}")]
            if not is_under_allowed_roots(docx_path, ALLOWED_ROOTS):
                logger.error(f"Path not allowed: {docx_path}")
                return [types.TextContent(type="text", text=f"Error: Path not allowed: {docx_path}")]
        # Convert to text
        logger.info(f"Converting DOCX to text: {docx_path}")
        text = extract_docx_text(docx_path)
        # Prefer writing to file if output_path provided
        if output_path_arg:
            out_path = Path(output_path_arg)
            if not is_under_allowed_roots(out_path, ALLOWED_ROOTS):
                logger.error(f"Output path not allowed: {out_path}")
                return [types.TextContent(type="text", text=f"Error: Output path not allowed: {out_path}")]
            out_path.parent.mkdir(parents=True, exist_ok=True)
            with open(out_path, "w", encoding="utf-8") as f:
                f.write(text)
            docx_size = docx_path.stat().st_size if docx_path.exists() else 0
            txt_size = out_path.stat().st_size
            msg = f"Success: {docx_size:,} bytes DOCX converted to {txt_size:,} bytes TXT at {out_path}"
            logger.info(msg)
            return [types.TextContent(type="text", text=msg)]
        # Otherwise return text to LLM
        logger.info("Returning extracted text to the LLM")
        return [types.TextContent(type="text", text=text)]
    except ImportError as e:
        logger.error(f"Dependency error: {e}")
        return [types.TextContent(type="text", text=f"Error: {str(e)}")]
    except httpx.HTTPError as e:
        logger.error(f"HTTP error downloading DOCX: {e}")
        return [types.TextContent(type="text", text=f"Error: HTTP error downloading DOCX: {str(e)}")]
    except Exception as e:
        logger.error(f"Error converting DOCX: {e}")
        return [types.TextContent(type="text", text=f"Error: {str(e)}")]
    finally:
        # Cleanup temporary file if used
        if temp_path:
            try:
                temp_path.unlink(missing_ok=True)
            except Exception:
                pass


# ============================================================================
# Tool Registration
# ============================================================================

@server.list_tools()
async def list_tools() -> list[types.Tool]:
    """List available tools"""
    return [
        types.Tool(
            name="convert_docx_to_text",
            description="Converts a DOCX file (local path or URL) to plain text; writes to output_path if provided, otherwise returns text.",
            inputSchema={
                "type": "object",
                "required": ["source"],
                "properties": {
                    "source": {
                        "type": "string",
                        "description": "Local .docx path or http/https URL"
                    },
                    "output_path": {
                        "type": "string",
                        "description": "Optional path to write the output text file"
                    },
                    "headers": {
                        "type": "object",
                        "description": "Optional HTTP headers for authenticated URL download (e.g., Authorization, Cookie)"
                    }
                }
            }
        )
    ]


# ============================================================================
# Tool Router
# ============================================================================

@server.call_tool()
async def tool_router(name: str, arguments: dict) -> list[types.TextContent]:
    """Route tool calls to appropriate handlers"""
    logger.info(f"Tool called: {name} with arguments: {arguments}")

    handlers = {
        "convert_docx_to_text": handle_convert_docx_to_text_tool
    }

    handler = handlers.get(name)
    if handler:
        return await handler(arguments)

    logger.error(f"Unknown tool: {name}")
    raise ValueError(f"Unknown tool: {name}")


# ============================================================================
# Resource Handlers
# ============================================================================

@server.list_resources()
async def list_resources() -> list[types.Resource]:
    return []


@server.list_resource_templates()
async def list_resource_templates() -> list[types.ResourceTemplate]:
    return []


# ============================================================================
# Create Starlette App
# ============================================================================

app = Starlette(
    debug=True,
    routes=[
        Route("/sse", endpoint=handle_sse, methods=["GET"]),
        Mount("/messages/", app=sse_transport.handle_post_message),
    ]
)


# ============================================================================
# Main Entry Point
# ============================================================================

if __name__ == "__main__":
    import uvicorn
    logger.info("Starting Converter MCP Server on http://0.0.0.0:8003")
    try:
        uvicorn.run(app, host="0.0.0.0", port=8003)
    except KeyboardInterrupt:
        logger.info("Server shutting down gracefully...")
        sys.exit(0)
    except Exception as e:
        logger.error(f"Server error: {e}")
        sys.exit(1)


"""
VSCode settings.json configuration:

{
  "mcpServers": {
    "convertermcp": {
      "type": "sse",
      "url": "http://localhost:8003/sse",
      "headers": {
        "Content-Type": "application/json"
      }
    }
  }
}
"""
